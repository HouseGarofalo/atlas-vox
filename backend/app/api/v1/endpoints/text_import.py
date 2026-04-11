"""Text import endpoints — extract text from URLs and files for synthesis."""

import re

import structlog
from fastapi import APIRouter, HTTPException, UploadFile
from pydantic import BaseModel, Field

from app.core.dependencies import CurrentUser

logger = structlog.get_logger(__name__)

router = APIRouter(prefix="/import-text", tags=["text-import"])

# Maximum upload size for imported text files (20 MB)
MAX_FILE_SIZE = 20 * 1024 * 1024


class URLImportRequest(BaseModel):
    url: str = Field(..., min_length=10, max_length=2000)


class TextImportResponse(BaseModel):
    text: str
    source: str
    character_count: int
    word_count: int


def _clean_text(text: str) -> str:
    """Clean extracted text — normalize whitespace, remove excessive newlines."""
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


@router.post("/url", response_model=TextImportResponse)
async def import_from_url(data: URLImportRequest, user: CurrentUser) -> TextImportResponse:
    """Extract readable text from a URL using readability heuristics."""
    logger.info("text_import_url", url=data.url)

    # Import SSRF protection from webhook dispatcher
    from app.services.webhook_dispatcher import _is_url_safe

    # SSRF protection: validate URL before fetching
    if not await _is_url_safe(data.url):
        raise HTTPException(
            status_code=400,
            detail="URL blocked: private/internal addresses are not allowed for security reasons"
        )

    try:
        import httpx
    except ImportError:
        raise HTTPException(status_code=501, detail="httpx not installed")

    try:
        async with httpx.AsyncClient(timeout=15, follow_redirects=False) as client:
            resp = await client.get(data.url, headers={"User-Agent": "Atlas-Vox-TextImport/1.0"})
            resp.raise_for_status()
            html = resp.text
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to fetch URL: {exc}")

    # Try readability extraction, fall back to simple HTML stripping
    text = ""
    try:
        from readability import Document
        doc = Document(html)
        # Get text from the readable content
        summary = doc.summary()
        text = re.sub(r"<[^>]+>", " ", summary)
    except ImportError:
        # Fallback: simple HTML tag stripping
        text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL)
        text = re.sub(r"<[^>]+>", " ", text)

    text = _clean_text(text)
    if not text:
        raise HTTPException(status_code=400, detail="Could not extract text from URL")

    # Limit to 50K chars
    if len(text) > 50000:
        text = text[:50000] + "\n\n[Truncated — original text exceeded 50,000 characters]"

    return TextImportResponse(
        text=text,
        source=data.url,
        character_count=len(text),
        word_count=len(text.split()),
    )


@router.post("/file", response_model=TextImportResponse)
async def import_from_file(file: UploadFile, user: CurrentUser) -> TextImportResponse:
    """Extract text from an uploaded file (TXT, PDF, DOCX)."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""

    content = await file.read(MAX_FILE_SIZE + 1)
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large. Maximum 20MB.")

    text = ""
    source = file.filename

    if ext == "txt":
        text = content.decode("utf-8", errors="replace")

    elif ext == "pdf":
        try:
            import io

            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages.append(page_text)
                text = "\n\n".join(pages)
        except ImportError:
            raise HTTPException(status_code=501, detail="pdfplumber not installed. PDF import unavailable.")

    elif ext == "docx":
        try:
            import io

            import docx
            doc = docx.Document(io.BytesIO(content))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise HTTPException(status_code=501, detail="python-docx not installed. DOCX import unavailable.")

    elif ext in ("epub", "html", "htm"):
        html_text = content.decode("utf-8", errors="replace")
        text = re.sub(r"<script[^>]*>.*?</script>", "", html_text, flags=re.DOTALL)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL)
        text = re.sub(r"<[^>]+>", " ", text)

    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: .{ext}. Supported: txt, pdf, docx, html")

    text = _clean_text(text)
    if not text:
        raise HTTPException(status_code=400, detail="Could not extract text from file")

    if len(text) > 50000:
        text = text[:50000] + "\n\n[Truncated]"

    return TextImportResponse(
        text=text,
        source=source,
        character_count=len(text),
        word_count=len(text.split()),
    )
